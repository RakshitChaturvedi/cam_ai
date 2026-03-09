from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any

def generate_cam_report(
    entity_name: str,
    majority_decision: Dict[str, Any],
    hitl_input: Dict[str, Any],
    output_filename: str = "Final_Credit_Appraisal_Memo.docx"
):
    """
    Generates a professional corporate Word Document summarizing the AI's transparent reasoning.
    """
    
    doc = Document()
    
    # 1. Header Block
    title = doc.add_heading(f"Credit Appraisal Memo (CAM)\nEntity: {entity_name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 2. Executive ML Decision Summary
    doc.add_heading("I. Executive Machine Learning Recommendation", level=1)
    
    # Highlight the AI decision in Bold
    p1 = doc.add_paragraph()
    runner = p1.add_run(f"Final Decision: {majority_decision.get('final_decision', 'N/A')}")
    runner.bold = True
    runner.font.size = Pt(14)
    # Give it a color depending on decision
    if majority_decision.get("final_decision") == "Approve":
        runner.font.color.rgb = RGBColor(0, 150, 0) # Green
    else:
        runner.font.color.rgb = RGBColor(200, 0, 0) # Red
        
    p2 = doc.add_paragraph()
    p2.add_run("ML Ensemble Confidence Score: ").bold = True
    p2.add_run(f"{majority_decision.get('ml_confidence_score', '0%')} ({majority_decision.get('winning_votes', '0/0')} Models Agreed)")
    
    p3 = doc.add_paragraph()
    p3.add_run("Recommended Limit: ").bold = True
    p3.add_run(f"₹ {majority_decision.get('recommended_limit_inr', 0):,}")
    
    p4 = doc.add_paragraph()
    p4.add_run("Risk Premium: ").bold = True
    p4.add_run(f"{majority_decision.get('risk_premium_percentage', 0.0)}%")

    p5 = doc.add_paragraph()
    p5.add_run("Primary Rationale: ").bold = True
    p5.add_run(str(majority_decision.get('primary_explainable_rationale', 'None provided.')))

    # 3. Transparent Human-In-The-Loop Adjustments
    doc.add_heading("II. Explainable HITL Adjustments", level=1)
    p6 = doc.add_paragraph("The Baseline Credit Score was deterministically adjusted based on Officer Field Notes:")
    
    if "adjustment_log" in hitl_input and hitl_input["adjustment_log"]:
        for log in hitl_input["adjustment_log"]:
            doc.add_paragraph(f"- {log}", style='List Bullet')
    else:
        doc.add_paragraph("No Human-in-the-Loop adjustments were applied.", style='Italic')
        
    # 4. Save
    doc.save(output_filename)
    print(f"\n[SUCCESS] Generated Corporate CAM: {output_filename}")
    return output_filename
